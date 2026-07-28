import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, List

class WordExporter:
    """
    Exports forensic reports to Microsoft Word format (.docx).
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        
    def export(self, report_data: Dict[str, Any], filename: str) -> str:
        """
        Generates a Word document from standard report data.
        """
        filepath = os.path.join(self.output_dir, filename)
        
        document = Document()
        
        # Add Title
        title = document.add_heading(report_data.get('title', 'Forensic Report'), 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        document.add_heading('Metadata', level=1)
        metadata = report_data.get('metadata', {})
        for key, value in metadata.items():
            document.add_paragraph(f"{key}: {value}")
            
        # Add sections
        sections = report_data.get('sections', [])
        for section in sections:
            document.add_heading(section.get('heading', 'Section'), level=1)
            document.add_paragraph(section.get('content', ''))
            
            # Add table if present
            if 'table_data' in section:
                table_data = section['table_data']
                if table_data:
                    headers = list(table_data[0].keys())
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = str(header)
                        
                    for row_data in table_data:
                        row_cells = table.add_row().cells
                        for i, header in enumerate(headers):
                            row_cells[i].text = str(row_data.get(header, ''))
                            
        document.save(filepath)
        return filepath
